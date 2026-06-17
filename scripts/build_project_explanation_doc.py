from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "internal_company_knowledge_assistant_explanation.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(89, 99, 110)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B7C9DD"


def set_run_font(run, *, name: str = "Calibri", size: float | None = None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color=None, bold=None):
    font = style.font
    font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        tag = f"w:{edge}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color="2E74B5", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", *, style=None, bold=False, color=None, size=None, italic=False, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_label_detail_table(doc, rows: list[tuple[str, str]], widths=(1700, 7660)):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, list(widths))
    set_table_borders(table)
    for row_idx, (label, detail) in enumerate(rows):
        label_cell, detail_cell = table.rows[row_idx].cells
        set_cell_shading(label_cell, HEADER_FILL)
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        run = label_p.add_run(label)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
        detail_p = detail_cell.paragraphs[0]
        detail_p.paragraph_format.space_after = Pt(0)
        run = detail_p.add_run(detail)
        set_run_font(run, size=9.5, color=INK)
    add_para(doc, "", after=4)
    return table


def add_matrix_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.3, color=DARK_BLUE, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for cell_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[cell_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, size=9, color=INK)
    add_para(doc, "", after=4)
    return table


def add_callout(doc, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D9E2EC")
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.7, color=INK)
    add_para(doc, "", after=4)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("Internal Company Knowledge Assistant")
    set_run_font(left, size=9, color=MUTED, bold=True)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    run = footer.add_run("Project explanation | FastAPI, Streamlit, LangChain, AWS")
    set_run_font(run, size=8.5, color=MUTED)
    return doc


def build() -> Path:
    OUT_DIR.mkdir(exist_ok=True)
    doc = setup_document()

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Internal Company Knowledge Assistant")
    set_run_font(run, size=24, color=INK, bold=True)
    subtitle = add_para(
        doc,
        "Detailed project explanation for the FastAPI, Streamlit, LangChain, Azure OpenAI, AWS, Langfuse, and RAGAS MVP.",
        size=12,
        color=MUTED,
        after=12,
    )
    subtitle.paragraph_format.keep_with_next = True
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="2E74B5", size="8")
    rule.paragraph_format.space_after = Pt(12)

    add_label_detail_table(
        doc,
        [
            ("Purpose", "Explain the implemented MVP architecture, code responsibilities, data flows, AWS services, evaluation strategy, and deployment path."),
            ("Audience", "Engineers, ML/LLM practitioners, cloud reviewers, and project assessors who need to understand how the assistant works."),
            ("Project state", "Greenfield scaffold implemented in this workspace with runnable backend/frontend containers and verification tests."),
            ("Primary outcome", "A chat-based internal knowledge assistant that retrieves grounded company answers from S3-indexed documents and records observability/evaluation metadata."),
        ],
    )

    add_heading(doc, "1. Executive Overview", 1)
    add_para(
        doc,
        "The project is an internal company knowledge assistant. Users sign in through a simple login screen, ask questions in a Streamlit chat interface, and receive grounded answers from a FastAPI backend. The backend coordinates a LangChain agent, an Azure OpenAI chat model, a RAG retriever over S3-indexed documents, persistent chat history, Langfuse tracing, and evaluation tooling.",
    )
    add_callout(
        doc,
        "MVP definition",
        "The MVP prioritizes a credible end-to-end architecture: secure secret loading from AWS Secrets Manager, document ingestion from S3, vector retrieval through OpenSearch Serverless, chat history persistence, transparent citations, token reporting, and repeatable eval/stress-test scripts.",
    )
    for item in [
        "FastAPI exposes authentication, health, chat, session history, and document catalog endpoints.",
        "Streamlit provides a chat-only user experience with login, previous sessions, response metadata, and citations.",
        "LangChain provides the agent abstraction and Azure OpenAI wrappers for chat and embeddings.",
        "AWS services provide object storage, container registry, serverless container hosting, secrets, logs, chat persistence, and vector search.",
        "Langfuse and RAGAS support observability, prompt versioning, golden-data scoring, and system-level evaluation.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. High-Level Architecture", 1)
    add_para(
        doc,
        "The assistant is split into a frontend service and a backend service. The frontend is intentionally thin: it handles login state, chat rendering, and API calls. The backend owns all sensitive operations, including secret loading, authentication, RAG retrieval, agent execution, chat history persistence, and tracing.",
    )
    add_matrix_table(
        doc,
        ["Layer", "Responsibility", "Main Services / Files"],
        [
            ["User interface", "Login, chat input, response display, session list, metadata panel.", "Streamlit; frontend/streamlit_app.py"],
            ["API layer", "Validates tokens, exposes chat/session/document endpoints, centralizes service wiring.", "FastAPI; backend/app/main.py"],
            ["Agent layer", "Builds context from tools and chat history, calls Azure OpenAI via LangChain.", "backend/app/agent.py"],
            ["Knowledge layer", "Reads S3 manifest, searches vectors, looks up CSV rows.", "S3, OpenSearch Serverless; storage.py, retrieval.py"],
            ["State layer", "Stores sessions and messages for persistent, contextual conversations.", "DynamoDB; history.py"],
            ["Observability/evals", "Prompt versioning, traces, token metadata, RAGAS reports, stress reports.", "Langfuse, RAGAS; observability.py, evals/"],
        ],
        [1700, 4100, 3560],
    )
    add_para(doc, "Runtime architecture:", bold=True, color=DARK_BLUE, after=3)
    add_para(
        doc,
        "User -> ALB -> Streamlit ECS service -> FastAPI ECS service -> Login/Secrets Manager, DynamoDB chat history, LangChain agent, OpenSearch RAG, S3 document catalog, Azure OpenAI, and Langfuse.",
        size=10.5,
    )

    add_heading(doc, "3. Backend Explanation", 1)
    add_para(
        doc,
        "The backend is the system of record for security, orchestration, and traceable answer generation. It keeps credentials out of the browser and uses dependency-style service factories so the application can switch between memory-backed local development and AWS-backed production services.",
    )
    add_label_detail_table(
        doc,
        [
            ("Configuration", "backend/app/config.py reads only non-secret environment values, such as region, secret names, S3 bucket, OpenSearch endpoint, and DynamoDB table."),
            ("Secrets", "backend/app/secrets.py loads JSON secrets from AWS Secrets Manager. Secret values are not read from .env files or committed config."),
            ("Authentication", "backend/app/auth.py verifies PBKDF2 password hashes and issues signed HMAC JWT-style bearer tokens."),
            ("Chat API", "backend/app/main.py exposes /auth/login, /chat, /chat/sessions, /chat/sessions/{id}, /documents, and /health."),
            ("Agent", "backend/app/agent.py loads prior history, runs tools, builds context, calls Azure OpenAI through LangChain, and returns traceable metadata."),
            ("Retries", "backend/app/retries.py wraps remote-facing paths with exponential backoff to reduce transient AWS/API failures."),
        ],
    )

    add_heading(doc, "4. Frontend Explanation", 1)
    add_para(
        doc,
        "The frontend is a Streamlit application optimized for a chat workflow. Users first see a login form. After successful authentication, the UI switches to a chat page, shows previous sessions in the sidebar, and renders assistant responses with response details.",
    )
    for item in [
        "The access token is held in Streamlit session state and sent as an Authorization bearer token.",
        "The chat input calls POST /chat with the current query and optional session_id.",
        "The sidebar can start a new chat, sign out, or reload older conversations.",
        "Each assistant response can display sources, tools used, input/output tokens, latency, and Langfuse trace ID.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Agentic AI Design", 1)
    add_para(
        doc,
        "The assistant uses an agentic pattern rather than a single prompt-only chain. The backend registers tools that expose different kinds of company knowledge, then the agent composes those tool results with the current query and recent chat history.",
    )
    add_matrix_table(
        doc,
        ["Tool", "Purpose", "When it helps"],
        [
            ["rag_search", "Semantic search over indexed document chunks with citation metadata.", "Policy, process, FAQ, handbook, or procedural questions."],
            ["document_catalog", "Lists and filters available indexed documents from the S3 manifest.", "Questions about what knowledge exists, document ownership, type, or department."],
            ["table_lookup", "Finds exact values from CSV-style sources in S3.", "Structured records such as contacts, escalation rows, owners, codes, or tabular policy values."],
        ],
        [1900, 3650, 3810],
    )
    add_callout(
        doc,
        "Why three tools matter",
        "RAG alone is useful, but not enough for a convincing agentic assistant. The catalog tool gives document awareness, while the table lookup tool improves exact-answer behavior for structured knowledge.",
    )

    add_heading(doc, "6. RAG And Ingestion Flow", 1)
    add_para(
        doc,
        "Documents are stored in S3 under a raw prefix. The ingestion job loads supported files, parses text, chunks content, generates embeddings with Azure OpenAI, indexes chunks into OpenSearch Serverless, and writes a document manifest back to S3.",
    )
    for item in [
        "Supported input formats: PDF, markdown, plain text, and CSV.",
        "Chunking uses LangChain text splitters when available, with a deterministic fallback splitter.",
        "Each chunk has a stable ID based on document key, chunk index, and checksum.",
        "The S3 manifest powers the document catalog tool and the /documents endpoint.",
        "The OpenSearch index stores text, title, URI, content type, chunk index, checksum, metadata, and optional embedding vector.",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "7. Chat History And Context", 1)
    add_para(
        doc,
        "Persistent chat history is central to follow-up questions. The backend stores both user and assistant messages. On every chat turn, it loads prior messages for the same user and session, builds a bounded context window, and passes that conversation history into the agent.",
    )
    add_label_detail_table(
        doc,
        [
            ("Local mode", "Uses an in-memory repository for fast development and unit testing."),
            ("Production mode", "Uses DynamoDB with user_id as the partition key and sort_key values for session summaries and messages."),
            ("Context control", "MAX_HISTORY_CHARS limits the amount of prior conversation injected into the prompt."),
            ("Long histories", "Older turns are omitted with a summary marker once the history window becomes too large."),
        ],
    )

    add_heading(doc, "8. Secrets And Security", 1)
    add_para(
        doc,
        "All secret values are expected to live in AWS Secrets Manager. Environment variables contain secret names and non-sensitive runtime configuration only. ECS task roles should receive least-privilege access to only the exact secret ARNs and AWS resources the service needs.",
    )
    add_matrix_table(
        doc,
        ["Secret", "Contains", "Used by"],
        [
            ["/company-assistant/{stage}/app", "session_secret and auth_users password-hash map.", "FastAPI authentication service."],
            ["/company-assistant/{stage}/azure-openai", "endpoint, api_key, api_version, chat deployment, embedding deployment.", "LangChain chat model and embedding model."],
            ["/company-assistant/{stage}/langfuse", "public key, secret key, base URL.", "Langfuse tracing and prompt management."],
        ],
        [2650, 4200, 2510],
    )
    add_callout(
        doc,
        "Security posture",
        "The implemented login is intentionally simple for the MVP. For production hardening, the natural next step is Cognito, SSO/SAML, or an internal identity provider while keeping the same backend authorization pattern.",
    )

    add_heading(doc, "9. Observability, Prompt Versioning, And Evals", 1)
    add_para(
        doc,
        "The project separates ordinary logging from LLM observability. CloudWatch captures container logs, while Langfuse captures traces, prompt versions, tool calls, latency, token usage, and failures for answer generation. RAGAS and stress testing provide system-level feedback.",
    )
    add_label_detail_table(
        doc,
        [
            ("Langfuse tracing", "Records model calls, retrieval/tool behavior, token usage, latency, and trace IDs that are returned to the UI."),
            ("Prompt versioning", "The system prompt can be loaded from Langfuse by prompt label, with a safe default prompt in code."),
            ("Golden-data eval", "evals/golden_dataset.csv defines question, expected answer, expected source, and tags."),
            ("RAGAS", "evals/run_ragas_eval.py can compute faithfulness, answer relevancy, context precision, and context recall when RAGAS dependencies and model access are configured."),
            ("Stress testing", "evals/stress_test.py generates 100 paraphrased queries and reports consistency, source overlap, latency, and failures."),
        ],
    )

    add_heading(doc, "10. AWS Deployment Explanation", 1)
    add_para(
        doc,
        "The deployment target is AWS ECS Fargate. Backend and frontend images are built separately, pushed to ECR, and deployed as services. An Application Load Balancer exposes the user-facing Streamlit service and can route API traffic internally or through a protected backend listener.",
    )
    add_matrix_table(
        doc,
        ["AWS service", "Role in the project", "Important configuration"],
        [
            ["S3", "Stores raw company documents and the ingestion manifest.", "Bucket policy, raw prefix, manifest key, encryption."],
            ["ECR", "Stores backend and frontend Docker images.", "Separate repositories and immutable tags for releases."],
            ["ECS Fargate", "Runs FastAPI and Streamlit containers.", "Task roles, service networking, CPU/memory, health checks."],
            ["ALB", "Provides browser access to the app.", "HTTPS listener, security groups, optional IP/VPN restriction."],
            ["Secrets Manager", "Stores all secret values.", "Least-privilege GetSecretValue permissions."],
            ["DynamoDB", "Persists chat sessions and messages.", "PAY_PER_REQUEST table with user_id/sort_key keys."],
            ["OpenSearch Serverless", "Stores vector index for RAG search.", "AOSS permissions, vector mapping, collection endpoint."],
            ["CloudWatch", "Captures ECS task logs.", "Backend and frontend log groups."],
        ],
        [1800, 3900, 3660],
    )

    add_heading(doc, "11. Repository Map", 1)
    add_matrix_table(
        doc,
        ["Path", "Explanation"],
        [
            ["backend/app/main.py", "FastAPI entrypoint, dependency wiring, auth enforcement, and API route definitions."],
            ["backend/app/agent.py", "Agent orchestration, context assembly, Azure OpenAI call, token estimates, and response metadata."],
            ["backend/app/ingest.py", "S3 document parsing, chunking, embedding, OpenSearch indexing, and manifest writing."],
            ["frontend/streamlit_app.py", "Login and chat interface, session sidebar, response details panel."],
            ["evals/", "Golden-data dataset, RAGAS runner, and 100-query stress-test runner."],
            ["infra/", "ECS task definitions, IAM policy, DynamoDB schema, and OpenSearch index mapping."],
            ["tests/", "Unit tests for auth, chat history, and agent/tool contract."],
        ],
        [2850, 6510],
    )

    add_heading(doc, "12. How A Query Is Answered", 1)
    for item in [
        "The user signs in through Streamlit, which calls POST /auth/login.",
        "FastAPI validates the username/password against password hashes loaded from AWS Secrets Manager.",
        "The user submits a chat question through Streamlit.",
        "FastAPI validates the bearer token and identifies the user.",
        "The backend loads previous messages for the session from memory or DynamoDB.",
        "The agent runs RAG search, document catalog lookup, and CSV/table lookup.",
        "The agent builds a prompt containing system instructions, chat history, tool context, and the user query.",
        "Azure OpenAI generates the answer through the LangChain wrapper.",
        "The backend stores the user and assistant messages, including sources, tools, token counts, latency, and trace ID.",
        "Streamlit displays the answer and lets the user inspect response metadata.",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "13. Verification Completed", 1)
    add_para(
        doc,
        "The implemented scaffold was verified without live cloud credentials. These checks prove the local code structure, core behavior, and deployment templates are internally consistent.",
    )
    add_label_detail_table(
        doc,
        [
            ("Python syntax", "compileall passed for backend, frontend, eval scripts, and tests."),
            ("Unit tests", "6 tests passed: authentication, token validation, password hashing, history persistence, bounded context, and agent tool contract."),
            ("Docker Compose", "docker compose config produced a valid service configuration."),
            ("Infra JSON", "ECS task definitions, IAM policy, DynamoDB table definition, and OpenSearch index mapping parsed successfully."),
        ],
    )

    add_heading(doc, "14. What Requires Live AWS/Azure Setup", 1)
    for item in [
        "Creating AWS Secrets Manager entries with real Azure OpenAI and Langfuse credentials.",
        "Creating the S3 bucket and uploading real company documents.",
        "Creating the OpenSearch Serverless collection and index with the expected vector dimension.",
        "Creating the DynamoDB chat history table.",
        "Building and pushing Docker images to ECR.",
        "Deploying ECS Fargate services and configuring the Application Load Balancer.",
        "Running ingestion, RAGAS evals, and the 100-query stress test against live infrastructure.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "15. Recommended Next Improvements", 1)
    add_matrix_table(
        doc,
        ["Improvement", "Why it matters", "Priority"],
        [
            ["Cognito or SSO", "Replaces simple login with enterprise identity and stronger lifecycle management.", "High"],
            ["Terraform or CDK", "Makes infrastructure repeatable and easier to review.", "High"],
            ["Admin ingestion UI", "Lets authorized users upload and re-index documents without CLI access.", "Medium"],
            ["Source preview", "Lets users inspect the exact retrieved snippets behind an answer.", "Medium"],
            ["Feedback buttons", "Captures human quality signals for future eval datasets.", "Medium"],
            ["PII redaction", "Reduces risk when indexing sensitive internal documents.", "Medium"],
            ["Cost dashboard", "Tracks token spend, retrieval latency, and per-user query volume.", "Low"],
        ],
        [2500, 5260, 1600],
    )

    add_heading(doc, "16. Final Summary", 1)
    add_para(
        doc,
        "This project is a strong five-day MVP because it demonstrates the full production shape of an internal knowledge assistant: containerized services, AWS hosting path, secure secret management, agentic tool use, RAG, persistent context, observability, prompt versioning, evaluation, and stress testing. The scaffold is intentionally extensible: it can run locally with development settings, then move toward production by configuring AWS resources and live Azure OpenAI/Langfuse secrets.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build()
    print(path)

