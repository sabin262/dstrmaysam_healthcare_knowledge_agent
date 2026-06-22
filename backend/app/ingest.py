from __future__ import annotations

import argparse
import hashlib
import io
import json
from dataclasses import dataclass
from typing import Any

from .aws import boto3_client, boto3_session
from .config import AppSettings
from .retries import retry_transient
from .secrets import SecretProvider


@dataclass
class ParsedDocument:
    key: str
    title: str
    text: str
    content_type: str
    checksum: str
    metadata: dict[str, Any]


def checksum_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def infer_healthcare_metadata(key: str, checksum: str) -> dict[str, Any]:
    normalized = key.lower()
    domain = "general"
    document_type = "document"
    allowed_roles = ["staff"]

    if any(marker in normalized for marker in ["clinical", "sop", "pathway", "guideline", "sepsis"]):
        domain = "clinical_policy"
        document_type = "policy"
        allowed_roles = ["doctor", "nurse", "clinical_governance", "admin"]
    elif any(marker in normalized for marker in ["hr", "training", "payroll", "onboarding"]):
        domain = "admin_policy"
        document_type = "policy"
        allowed_roles = ["staff", "admin", "manager"]
    elif any(marker in normalized for marker in ["incident", "breach", "safeguarding", "governance"]):
        domain = "compliance"
        document_type = "policy"
        allowed_roles = ["staff", "manager", "clinical_governance"]
    elif any(marker in normalized for marker in ["catalogue", "directory", "service", "owner", "system"]):
        domain = "catalogue"
        document_type = "directory"
    elif any(marker in normalized for marker in ["calendar", "rota", "on-call", "oncall", "clinic"]):
        domain = "rota"
        document_type = "schedule"
    elif any(marker in normalized for marker in ["formulary", "medicine", "drug", "restricted"]):
        domain = "formulary"
        document_type = "table"
        allowed_roles = ["doctor", "nurse", "pharmacy", "clinical_governance", "admin"]

    return {
        "key": key,
        "checksum": checksum,
        "owner": "unknown",
        "version": "unknown",
        "effective_date": "unknown",
        "review_date": "unknown",
        "approval_status": "unknown",
        "sensitivity": "internal",
        "domain": domain,
        "document_type": document_type,
        "allowed_roles": allowed_roles,
    }


def parse_document(key: str, data: bytes) -> ParsedDocument:
    lower = key.lower()
    title = key.rsplit("/", 1)[-1]
    checksum = checksum_bytes(data)
    metadata = infer_healthcare_metadata(key, checksum)

    if lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            text = f"PDF parsing failed for {key}: {exc}"
        content_type = "application/pdf"
    elif lower.endswith(".docx"):
        try:
            from docx import Document

            document = Document(io.BytesIO(data))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            table_rows = []
            for table in document.tables:
                for row in table.rows:
                    table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n".join(paragraphs + table_rows)
        except Exception as exc:
            text = f"DOCX parsing failed for {key}: {exc}"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif lower.endswith(".csv"):
        text = data.decode("utf-8", errors="replace")
        content_type = "text/csv"
    elif lower.endswith(".md"):
        text = data.decode("utf-8", errors="replace")
        content_type = "text/markdown"
    else:
        text = data.decode("utf-8", errors="replace")
        content_type = "text/plain"

    return ParsedDocument(
        key=key,
        title=title,
        text=text,
        content_type=content_type,
        checksum=checksum,
        metadata=metadata,
    )


def chunk_text(text: str, chunk_size: int = 1500, chunk_overlap: int = 250) -> list[str]:
    chunk_size = max(300, chunk_size)
    chunk_overlap = max(0, min(chunk_overlap, chunk_size - 1))
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return splitter.split_text(text)
    except Exception:
        chunks: list[str] = []
        step = chunk_size
        overlap = chunk_overlap
        index = 0
        while index < len(text):
            chunks.append(text[index : index + step])
            index += step - overlap
    return chunks


def is_metadata_only_manifest_record(record: dict[str, Any]) -> bool:
    metadata = record.get("metadata") if isinstance(record.get("metadata"), dict) else {}
    return (
        str(record.get("ingestion_status") or "") == "metadata_only"
        or str(metadata.get("asset_source") or "") == "postgres_uploaded_lookup"
        or str(record.get("uri") or "").startswith("postgres://")
        or str(record.get("key") or "").startswith("postgres://")
    )


class IngestionJob:
    def __init__(self, settings: AppSettings, secret_provider: SecretProvider):
        self.settings = settings
        self.secret_provider = secret_provider
        self.s3 = boto3_client(settings, "s3")
        self._embeddings: Any | None = None
        self._opensearch: Any | None = None

    @retry_transient
    def run(self) -> dict[str, Any]:
        existing_manifest = self._load_existing_manifest()
        previous_index = existing_manifest.get("opensearch_index")
        force_reindex = previous_index != self.settings.opensearch_index
        existing_by_key = {
            str(document.get("key", "")): document
            for document in existing_manifest.get("documents", [])
            if isinstance(document, dict) and document.get("key") and not is_metadata_only_manifest_record(document)
        }
        metadata_only_documents = [
            dict(document)
            for document in existing_manifest.get("documents", [])
            if isinstance(document, dict) and is_metadata_only_manifest_record(document)
        ]
        raw_documents = self._load_raw_documents()
        seen_keys: set[str] = set()
        indexed_chunks = 0
        indexed_documents = 0
        skipped_documents = 0
        deleted_documents = 0
        deleted_chunks = 0
        manifest_documents = list(metadata_only_documents)
        for raw_document in raw_documents:
            key = raw_document["key"]
            seen_keys.add(key)
            checksum = checksum_bytes(raw_document["body"])
            existing_document = existing_by_key.get(key)
            if existing_document and existing_document.get("checksum") == checksum and not force_reindex:
                skipped_documents += 1
                unchanged_document = dict(existing_document)
                unchanged_document["ingestion_status"] = "skipped_unchanged"
                manifest_documents.append(unchanged_document)
                continue

            if existing_document and not force_reindex:
                deleted_chunks += self._delete_document_chunks(key)

            document = parse_document(key, raw_document["body"])
            chunks = chunk_text(
                document.text,
                chunk_size=self.settings.ingestion_chunk_size,
                chunk_overlap=self.settings.ingestion_chunk_overlap,
            )
            for chunk_index, chunk in enumerate(chunks):
                self._index_chunk(document, chunk, chunk_index)
                indexed_chunks += 1
            indexed_documents += 1
            manifest_documents.append(
                {
                    "key": document.key,
                    "title": document.title,
                    "content_type": document.content_type,
                    "checksum": document.checksum,
                    "metadata": document.metadata,
                    "chunk_count": len(chunks),
                    "ingestion_status": "indexed",
                }
            )

        removed_keys = sorted(set(existing_by_key) - seen_keys)
        if not force_reindex:
            for key in removed_keys:
                deleted_chunks += self._delete_document_chunks(key)
                deleted_documents += 1

        total_chunks = sum(int(document.get("chunk_count") or 0) for document in manifest_documents)
        manifest = {
            "opensearch_index": self.settings.opensearch_index,
            "previous_opensearch_index": previous_index,
            "force_reindex": force_reindex,
            "documents": manifest_documents,
            "indexed_chunks": indexed_chunks,
            "total_chunks": total_chunks,
            "indexed_documents": indexed_documents,
            "skipped_documents": skipped_documents,
            "deleted_documents": deleted_documents,
            "deleted_chunks": deleted_chunks,
        }
        self.s3.put_object(
            Bucket=self.settings.s3_bucket,
            Key=self.settings.s3_manifest_key,
            Body=json.dumps(manifest, indent=2).encode("utf-8"),
            ContentType="application/json",
        )
        return manifest

    def _load_existing_manifest(self) -> dict[str, Any]:
        try:
            response = self.s3.get_object(
                Bucket=self.settings.s3_bucket,
                Key=self.settings.s3_manifest_key,
            )
            manifest = json.loads(response["Body"].read().decode("utf-8"))
            return manifest if isinstance(manifest, dict) else {"documents": []}
        except Exception:
            return {"documents": []}

    def _load_raw_documents(self) -> list[dict[str, Any]]:
        paginator = self.s3.get_paginator("list_objects_v2")
        documents: list[dict[str, Any]] = []
        for page in paginator.paginate(Bucket=self.settings.s3_bucket, Prefix=self.settings.s3_raw_prefix):
            for item in page.get("Contents", []):
                key = item["Key"]
                if key.endswith("/"):
                    continue
                if not key.lower().endswith((".pdf", ".docx", ".txt", ".md")):
                    continue
                body = self.s3.get_object(Bucket=self.settings.s3_bucket, Key=key)["Body"].read()
                documents.append({"key": key, "body": body})
        return documents

    def _load_documents(self) -> list[ParsedDocument]:
        return [parse_document(document["key"], document["body"]) for document in self._load_raw_documents()]

    def _index_chunk(self, document: ParsedDocument, chunk: str, chunk_index: int) -> None:
        client = self._get_opensearch()
        embedding = self._embed(chunk)
        doc_id = hashlib.sha256(f"{document.key}:{chunk_index}:{document.checksum}".encode()).hexdigest()
        body = {
            "key": document.key,
            "title": document.title,
            "uri": f"s3://{self.settings.s3_bucket}/{document.key}",
            "text": chunk,
            "content_type": document.content_type,
            "chunk_index": chunk_index,
            "checksum": document.checksum,
            "metadata": document.metadata,
        }
        if embedding is not None:
            body["embedding"] = embedding
        client.index(index=self.settings.opensearch_index, id=doc_id, body=body, refresh=False)

    def _delete_document_chunks(self, key: str) -> int:
        client = self._get_opensearch()
        try:
            response = client.delete_by_query(
                index=self.settings.opensearch_index,
                body={"query": {"term": {"key": key}}},
                refresh=True,
                conflicts="proceed",
            )
            return int(response.get("deleted") or 0)
        except Exception:
            return 0

    def _embed(self, text: str) -> list[float] | None:
        try:
            if self._embeddings is None:
                from langchain_openai import AzureOpenAIEmbeddings

                secrets = self.secret_provider.load_azure_openai()
                self._embeddings = AzureOpenAIEmbeddings(
                    azure_endpoint=secrets.endpoint,
                    api_key=secrets.api_key,
                    api_version=secrets.api_version,
                    azure_deployment=secrets.embedding_deployment,
                )
            return list(self._embeddings.embed_query(text))
        except Exception:
            return None

    def _get_opensearch(self) -> Any:
        if self._opensearch is not None:
            return self._opensearch
        from opensearchpy import OpenSearch, RequestsHttpConnection
        from opensearchpy import AWSV4SignerAuth

        if not self.settings.opensearch_endpoint:
            raise RuntimeError("OPENSEARCH_ENDPOINT is required for ingestion")
        credentials = boto3_session(self.settings).get_credentials()
        auth = AWSV4SignerAuth(credentials, self.settings.aws_region, "aoss")
        host = self.settings.opensearch_endpoint.replace("https://", "").replace("http://", "")
        self._opensearch = OpenSearch(
            hosts=[{"host": host, "port": 443}],
            http_auth=auth,
            use_ssl=True,
            verify_certs=True,
            connection_class=RequestsHttpConnection,
        )
        return self._opensearch


def main() -> int:
    parser = argparse.ArgumentParser(description="Ingest S3 documents into OpenSearch")
    parser.parse_args()
    settings = AppSettings.from_env()
    result = IngestionJob(settings, SecretProvider(settings)).run()
    print(json.dumps(result, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
